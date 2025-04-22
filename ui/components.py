import streamlit as st
import base64
from io import BytesIO
from fpdf import FPDF
from docx import Document
from docx.shared import Pt, Inches
import plotly.express as px
import plotly.graph_objects as go
import pandas as pd
import json
import io

def custom_styles():
    """Apply custom CSS styling to the application"""
    # Main styling
    st.markdown("""
    <style>
        /* Main container styling */
        .main .block-container {
            padding-top: 2rem;
            padding-bottom: 2rem;
        }
        
        /* Header styling */
        h1 {
            color: #2e3c54;
            font-weight: 600;
            margin-bottom: 1.5rem;
        }
        
        h2 {
            color: #3c4c64;
            font-weight: 500;
            margin-top: 1.5rem;
            margin-bottom: 1rem;
        }
        
        h3 {
            color: #475b74;
            font-weight: 500;
            margin-top: 1.2rem;
            margin-bottom: 0.8rem;
        }
        
        /* Card and container styling */
        .info-box {
            background-color: #f0f7ff;
            padding: 1rem;
            border-radius: 0.5rem;
            margin-bottom: 1rem;
            border-left: 4px solid #4a86e8;
        }
        
        .resource-card {
            background-color: white;
            padding: 1rem;
            border-radius: 0.5rem;
            margin-bottom: 0.8rem;
            border: 1px solid #e0e0e0;
            box-shadow: 0 2px 4px rgba(0,0,0,0.05);
        }
        
        .results-container {
            background-color: #f9f9f9;
            padding: 1.5rem;
            border-radius: 0.5rem;
            margin-top: 1rem;
            border-left: 5px solid #6495ED;
        }
        
        /* Button styling */
        .stButton button {
            border-radius: 4px;
            transition: all 0.3s;
        }
        
        .stButton button:hover {
            transform: translateY(-1px);
            box-shadow: 0 4px 8px rgba(0,0,0,0.1);
        }
        
        /* Tabs styling */
        .stTabs [data-baseweb="tab-list"] {
            gap: 8px;
        }
        
        .stTabs [data-baseweb="tab"] {
            height: 40px;
            border-radius: 4px 4px 0px 0px;
            padding-left: 16px;
            padding-right: 16px;
        }
        
        /* Table styling */
        .stDataFrame {
            border-radius: 8px;
            overflow: hidden;
        }
        
        /* Custom element styling */
        .char-counter {
            color: #6a6a6a;
            font-size: 0.8rem;
            text-align: right;
        }
        
        .progress-container {
            margin: 1rem 0;
        }
        
        /* Icon styling */
        .icon-text {
            display: flex;
            align-items: center;
        }
        
        .icon-text svg {
            margin-right: 0.5rem;
        }
    </style>
    """, unsafe_allow_html=True)

def info_box(title, content, icon="ℹ️"):
    """Display an information box with title and content"""
    html = f"""
    <div class="info-box">
        <h4 style="margin-top: 0; margin-bottom: 0.5rem; display: flex; align-items: center;">
            <span style="margin-right: 8px;">{icon}</span> {title}
        </h4>
        <p style="margin-bottom: 0;">{content}</p>
    </div>
    """
    st.markdown(html, unsafe_allow_html=True)

def resource_card(title, description, link=None, icon="📚"):
    """Display a resource card with title, description, and optional link"""
    html = f"""
    <div class="resource-card">
        <h4 style="margin-top: 0; margin-bottom: 0.5rem; display: flex; align-items: center;">
            <span style="margin-right: 8px;">{icon}</span> {title}
        </h4>
        <p style="margin-bottom: {0 if link else '0'};">{description}</p>
    """
    
    if link:
        html += f'<a href="{link}" target="_blank" style="text-decoration: none;">Learn More →</a>'
    
    html += "</div>"
    st.markdown(html, unsafe_allow_html=True)

def progress_step(title, description, is_active=False, is_completed=False):
    """Display a progress step indicator"""
    # Determine the status icon
    icon = "⚪"  # Default (not started)
    if is_completed:
        icon = "✅"  # Completed
    elif is_active:
        icon = "🔵"  # Active
    
    # Determine text color
    color = "#6a6a6a"  # Default (not started)
    if is_completed:
        color = "#2e7d32"  # Completed (green)
    elif is_active:
        color = "#1976d2"  # Active (blue)
    
    html = f"""
    <div style="display: flex; margin-bottom: 1rem; align-items: flex-start;">
        <div style="margin-right: 10px; font-size: 1.2rem;">{icon}</div>
        <div>
            <h4 style="margin: 0; color: {color};">{title}</h4>
            <p style="margin: 0; color: #6a6a6a; font-size: 0.9rem;">{description}</p>
        </div>
    </div>
    """
    st.markdown(html, unsafe_allow_html=True)

def display_progress_steps(current_step):
    """Display the complete progress steps for the assessment process"""
    steps = [
        {
            "title": "Begin Assessment",
            "description": "Share your grief experience and answer reflective questions"
        },
        {
            "title": "Review Results",
            "description": "Get personalized insights based on your responses"
        },
        {
            "title": "Receive Guidance",
            "description": "Get a tailored guide with coping strategies and resources"
        },
        {
            "title": "Track Progress",
            "description": "Save your assessment and track changes over time"
        }
    ]
    
    # Map current_step to index
    step_map = {
        "intro": -1,  # Before starting
        "assessment": 0,
        "processing": 0,
        "results": 1,
        "guide_selection": 2,
        "guide_view": 2,
        "history": 3
    }
    
    current_index = step_map.get(current_step, -1)
    
    for i, step in enumerate(steps):
        is_active = (i == current_index)
        is_completed = (i < current_index)
        progress_step(step["title"], step["description"], is_active, is_completed)

def chat_message_user(message):
    """Display a user message in chat format"""
    with st.chat_message("user"):
        st.markdown(message)

def chat_message_ai(message):
    """Display an AI message in chat format"""
    with st.chat_message("assistant"):
        st.markdown(message)

def render_header(title, logo_path=None, use_container_width=True):
    """
    Render the header section with logo and title
    
    Args:
        title (str): The title to display
        logo_path (str, optional): Path to the logo image
        use_container_width (bool): Whether to use the full container width
    """
    col1, col2 = st.columns([1, 3])
    
    try:
        with col1:
            if logo_path:
                st.image(logo_path, width=100)
            else:
                # Fallback to an icon or text
                st.markdown("🧭")
        
        with col2:
            st.title(title)
    except Exception as e:
        st.error(f"Error rendering header: {str(e)}")
        # Fallback to simple header
        st.title(title)

def render_step_progress(current_step, total_steps):
    """
    Render a progress bar for the current step in the assessment
    
    Args:
        current_step (int): The current step number
        total_steps (int): The total number of steps
    """
    progress = current_step / total_steps
    st.progress(progress)
    st.caption(f"Step {current_step} of {total_steps}")

def render_navigation_buttons(back_action=None, next_action=None, back_text="Back", next_text="Next", custom_buttons=None):
    """
    Render navigation buttons for moving between steps
    
    Args:
        back_action (callable, optional): Action to perform when back button is clicked
        next_action (callable, optional): Action to perform when next button is clicked
        back_text (str): Text to display on the back button
        next_text (str): Text to display on the next button
        custom_buttons (list, optional): List of dicts with 'text', 'action', and 'primary' keys
    """
    cols = st.columns([1, 1, 1])
    
    with cols[0]:
        if back_action:
            if st.button(back_text, key="back_button", use_container_width=True):
                back_action()
    
    with cols[1]:
        if custom_buttons:
            for button in custom_buttons:
                if button.get('position', 1) == 1:
                    if st.button(
                        button['text'], 
                        key=f"custom_button_{button['text']}", 
                        type="primary" if button.get('primary', False) else "secondary",
                        use_container_width=True
                    ):
                        button['action']()
    
    with cols[2]:
        if next_action:
            if st.button(next_text, key="next_button", type="primary", use_container_width=True):
                next_action()
        elif custom_buttons:
            for button in custom_buttons:
                if button.get('position', 1) == 2:
                    if st.button(
                        button['text'], 
                        key=f"custom_button_{button['text']}", 
                        type="primary" if button.get('primary', False) else "secondary",
                        use_container_width=True
                    ):
                        button['action']()

def render_fixed_navigation():
    """Render a fixed navigation button in the bottom right corner"""
    st.markdown(
        """
        <style>
        .fixed-nav {
            position: fixed;
            right: 30px;
            bottom: 30px;
            width: 120px;
            z-index: 1000;
        }
        </style>
        """,
        unsafe_allow_html=True
    )
    
    # Create a container for the fixed button
    with st.container():
        st.markdown('<div class="fixed-nav">', unsafe_allow_html=True)
        st.button("Next", key="fixed_next", type="primary", use_container_width=True)
        st.markdown('</div>', unsafe_allow_html=True)

def create_emotion_visualization(emotions, intensities):
    """
    Create a bar chart visualization of emotions
    
    Args:
        emotions (list): List of emotion labels
        intensities (list): List of intensity values for each emotion
        
    Returns:
        plotly figure for bar chart
    """
    if not emotions or not intensities or len(emotions) != len(intensities):
        return None
    
    # Create dataframe for plotting
    df = pd.DataFrame({
        'Emotion': emotions,
        'Intensity': intensities
    })
    
    # Sort by intensity for the bar chart
    df_sorted = df.sort_values('Intensity', ascending=False)
    
    # Create bar chart
    fig_bar = px.bar(
        df_sorted, 
        x='Emotion', 
        y='Intensity',
        color='Intensity',
        color_continuous_scale='Viridis',
        title="Emotion Intensity",
        height=400
    )
    
    fig_bar.update_layout(
        xaxis_title="",
        yaxis_title="Intensity (0-10)",
        coloraxis_showscale=False
    )
    
    return fig_bar

def render_emotion_display(responses):
    """
    Render a visual display of emotions based on assessment responses
    
    Args:
        responses (dict): The assessment responses
    """
    # Extract emotion data from responses
    emotions_response = responses.get("emotions", {}).get("response", [])
    if not emotions_response:
        st.info("No emotion data available to visualize.")
        return
    
    # Get intensity values
    emotion_intensity = responses.get("emotion_intensity", {}).get("response", 5)
    
    # Convert to numeric values if needed
    if isinstance(emotion_intensity, str):
        try:
            emotion_intensity = float(emotion_intensity)
        except:
            emotion_intensity = 5
    
    # For demonstration, let's create simulated intensity values
    # In real use, these would come from the assessment
    from config.questions import get_emotion_labels
    
    all_emotions = get_emotion_labels()
    selected_emotions = []
    intensities = []
    
    # Create intensity values for selected emotions
    for emotion in all_emotions:
        if emotion in emotions_response:
            selected_emotions.append(emotion)
            # Either use provided intensity or a random value between 5-9
            import random
            intensities.append(emotion_intensity if len(emotions_response) == 1 else random.randint(5, 9))
    
    if not selected_emotions:
        st.info("No specific emotions were selected in the assessment.")
        return
    
    # Create visualization
    fig_bar = create_emotion_visualization(selected_emotions, intensities)
    
    if fig_bar:
        st.plotly_chart(fig_bar, use_container_width=True)
        st.caption("This bar chart ranks your emotions by intensity.")
    else:
        st.info("Unable to create emotion visualization with the provided data.")

def render_guide_download_options(guide_data):
    """
    Render download buttons for the guide in various formats
    
    Args:
        guide_data (dict): The guide data to convert to downloadable formats
    """
    st.markdown("### Download Your Personalized Guide")
    
    # Create columns for the download buttons
    col1, col2, col3 = st.columns(3)
    
    with col1:
        pdf_bytes = create_guide_pdf(guide_data)
        if pdf_bytes:
            st.download_button(
                label="Download as PDF",
                data=pdf_bytes,
                file_name="Grief_Compass_Guide.pdf",
                mime="application/pdf",
                key="download_pdf",
                use_container_width=True
            )
    
    with col2:
        docx_bytes = create_guide_docx(guide_data)
        if docx_bytes:
            st.download_button(
                label="Download as DOCX",
                data=docx_bytes,
                file_name="Grief_Compass_Guide.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                key="download_docx",
                use_container_width=True
            )
    
    with col3:
        txt_bytes = create_guide_txt(guide_data)
        if txt_bytes:
            st.download_button(
                label="Download as TXT",
                data=txt_bytes,
                file_name="Grief_Compass_Guide.txt",
                mime="text/plain",
                key="download_txt",
                use_container_width=True
            )

def create_guide_pdf(guide_data):
    """
    Create a PDF version of the grief guide.
    
    Args:
        guide_data (dict): The guide data to include in the PDF
        
    Returns:
        bytes: PDF file as bytes
    """
    try:
        from fpdf import FPDF
        
        class CustomPDF(FPDF):
            def header(self):
                # Header on each page
                self.set_font('Arial', 'B', 12)
                self.cell(0, 10, 'Grief Compass - Your Personalized Guide', 0, 1, 'C')
                self.ln(5)
            
            def footer(self):
                # Footer on each page
                self.set_y(-15)
                self.set_font('Arial', 'I', 8)
                self.cell(0, 10, f'Page {self.page_no()}/{{nb}}', 0, 0, 'C')
        
        # Create PDF object with UTF-8 support
        pdf = CustomPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.alias_nb_pages()
        pdf.add_page()
        # Use standard Arial font which is built-in
        pdf.set_font('Arial', '', 12)
        
        # Helper function to replace problematic characters
        def clean_text(text):
            if not text:
                return ""
            # Replace bullet points and other problematic characters
            text = text.replace('•', '-').replace('\u2022', '-')
            # Replace other potentially problematic Unicode characters
            text = ''.join(c if ord(c) < 128 else '-' for c in text)
            return text
        
        # Title
        pdf.set_font('Arial', 'B', 16)
        pdf.cell(0, 10, clean_text(guide_data.get('title', 'Your Personalized Grief Guide')), 0, 1, 'C')
        pdf.ln(5)
        
        # Introduction
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, 'Introduction', 0, 1)
        pdf.set_font('Arial', '', 12)
        pdf.multi_cell(0, 10, clean_text(guide_data.get('introduction', '')))
        pdf.ln(5)
        
        # Include responses and results summary if available from session state
        if 'responses' in st.session_state and 'results' in st.session_state:
            # Add key information about the loss
            pdf.set_font('Arial', 'B', 14)
            pdf.cell(0, 10, 'About Your Loss', 0, 1)
            pdf.set_font('Arial', '', 12)
            
            cause = st.session_state.responses.get("cause_of_death", {}).get("response", "Not specified")
            time_period = st.session_state.responses.get("time_since_loss", {}).get("response", "Not specified")
            relationship = st.session_state.responses.get("relationship", {}).get("response", "Not specified")
            
            pdf.cell(0, 10, f"Cause: {clean_text(cause)}", 0, 1)
            pdf.cell(0, 10, f"Relationship: {clean_text(relationship)}", 0, 1)
            pdf.cell(0, 10, f"Time Since Loss: {clean_text(time_period)}", 0, 1)
            pdf.ln(5)
            
            # Add assessment summary
            if st.session_state.results and "summary" in st.session_state.results:
                pdf.set_font('Arial', 'B', 14)
                pdf.cell(0, 10, 'Assessment Summary', 0, 1)
                pdf.set_font('Arial', '', 12)
                pdf.multi_cell(0, 10, clean_text(st.session_state.results.get("summary", "")))
                pdf.ln(5)
            
            # Add recommendations
            if st.session_state.results and "recommendations" in st.session_state.results:
                pdf.set_font('Arial', 'B', 14)
                pdf.cell(0, 10, 'Recommendations', 0, 1)
                pdf.set_font('Arial', '', 12)
                
                for rec in st.session_state.results.get("recommendations", []):
                    pdf.cell(10)  # Indent
                    pdf.multi_cell(0, 10, f"- {clean_text(rec)}")
                pdf.ln(5)
        
        # Weekly Routine
        pdf.add_page()
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, 'Your 7-Day Healing Plan', 0, 1)
        pdf.set_font('Arial', '', 12)
        
        # Display daily routines
        daily_routines = guide_data.get("weekly_routine", {})
        
        for day, routines in daily_routines.items():
            pdf.set_font('Arial', 'B', 12)
            pdf.cell(0, 10, f"Day {day}: {clean_text(routines.get('key_focus', ''))}", 0, 1)
            pdf.set_font('Arial', '', 11)
            
            if "hourly_schedule" in routines:
                hourly_schedule = routines["hourly_schedule"]
                for hour, activity in hourly_schedule.items():
                    pdf.set_font('Arial', 'B', 10)
                    pdf.cell(20, 10, f"{hour}:", 0, 0)
                    pdf.set_font('Arial', '', 10)
                    pdf.multi_cell(0, 10, clean_text(activity))
            pdf.ln(5)
        
        # Self-Care
        pdf.add_page()
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, 'Self-Care Strategies', 0, 1)
        pdf.set_font('Arial', '', 12)
        
        self_care = guide_data.get("self_care", {})
        
        # Physical activity
        pdf.set_font('Arial', 'B', 12)
        pdf.cell(0, 10, 'Physical Activity', 0, 1)
        pdf.set_font('Arial', '', 11)
        pdf.multi_cell(0, 10, clean_text(self_care.get("physical_activity", "")))
        pdf.ln(5)
        
        # Nourishing meal
        pdf.set_font('Arial', 'B', 12)
        pdf.cell(0, 10, 'Nourishing Meal', 0, 1)
        pdf.set_font('Arial', '', 11)
        pdf.multi_cell(0, 10, clean_text(self_care.get("nourishing_meal", "")))
        pdf.ln(5)
        
        # Evening ritual
        pdf.set_font('Arial', 'B', 12)
        pdf.cell(0, 10, 'Evening Ritual', 0, 1)
        pdf.set_font('Arial', '', 11)
        pdf.multi_cell(0, 10, clean_text(self_care.get("evening_ritual", "")))
        pdf.ln(5)
        
        # Resources
        pdf.add_page()
        pdf.set_font('Arial', 'B', 14)
        pdf.cell(0, 10, 'Resources', 0, 1)
        pdf.set_font('Arial', '', 12)
        
        resources = guide_data.get("resources", {})
        
        # Support groups
        if "support_groups" in resources and resources["support_groups"]:
            pdf.set_font('Arial', 'B', 12)
            pdf.cell(0, 10, 'Support Groups', 0, 1)
            pdf.set_font('Arial', '', 11)
            
            for group in resources["support_groups"]:
                pdf.set_font('Arial', 'B', 10)
                pdf.cell(0, 10, clean_text(group.get("name", "")), 0, 1)
                pdf.set_font('Arial', '', 10)
                pdf.cell(0, 5, f"Website: {clean_text(group.get('url', ''))}", 0, 1)
                pdf.multi_cell(0, 10, clean_text(group.get("description", "")))
                pdf.ln(5)
        
        # Books
        if "books" in resources and resources["books"]:
            pdf.set_font('Arial', 'B', 12)
            pdf.cell(0, 10, 'Books', 0, 1)
            pdf.set_font('Arial', '', 11)
            
            for book in resources["books"]:
                pdf.set_font('Arial', 'B', 10)
                pdf.cell(0, 10, f"{clean_text(book.get('title', ''))} by {clean_text(book.get('author', ''))}", 0, 1)
                pdf.set_font('Arial', '', 10)
                pdf.multi_cell(0, 10, clean_text(book.get("description", "")))
                pdf.ln(5)
        
        # Hotlines
        if "hotlines" in resources and resources["hotlines"]:
            pdf.set_font('Arial', 'B', 12)
            pdf.cell(0, 10, 'Hotlines & Crisis Support', 0, 1)
            pdf.set_font('Arial', '', 11)
            
            for hotline in resources["hotlines"]:
                pdf.set_font('Arial', 'B', 10)
                pdf.cell(0, 10, clean_text(hotline.get("name", "")), 0, 1)
                pdf.set_font('Arial', '', 10)
                pdf.cell(0, 5, f"Number: {clean_text(hotline.get('number', ''))}", 0, 1)
                pdf.multi_cell(0, 10, clean_text(hotline.get("description", "")))
                pdf.ln(5)
        
        # Output PDF
        try:
            # Try to output directly as bytes
            output = pdf.output(dest='S')
            if isinstance(output, bytes):
                return output
            else:
                # For older FPDF versions that return strings
                return output.encode('latin-1', errors='replace')
        except Exception as e:
            # Alternative approach for newer FPDF versions
            import io
            output_buffer = io.BytesIO()
            pdf.output(output_buffer)
            output_buffer.seek(0)
            return output_buffer.getvalue()
        
    except Exception as e:
        st.error(f"Error creating PDF: {str(e)}")
        return None

def create_guide_docx(guide):
    """Create a DOCX version of the personalized guide"""
    try:
        # Create document
        doc = Document()
        
        # Add title
        doc.add_heading(guide.get('title', 'Your Personalized Grief Support Guide'), 0)
        
        # Add introduction
        doc.add_heading('Overview', level=1)
        doc.add_paragraph(guide.get('introduction', ''))
        
        # Weekly Routine Section
        doc.add_heading('Your 7-Day Healing Plan', level=1)
        doc.add_paragraph('This weekly routine is designed specifically for your grief journey.')
        
        daily_routines = guide.get('weekly_routine', {})
        for day, routines in daily_routines.items():
            doc.add_heading(f"Day {day}: {routines.get('key_focus', '')}", level=2)
            
            if "hourly_schedule" in routines:
                hourly_schedule = routines["hourly_schedule"]
                for hour, activity in hourly_schedule.items():
                    p = doc.add_paragraph()
                    p.add_run(f"{hour}:").bold = True
                    doc.add_paragraph(activity)
            else:
                # Fallback for old format
                if 'morning' in routines:
                    p = doc.add_paragraph()
                    p.add_run('Morning:').bold = True
                    doc.add_paragraph(routines['morning'])
                
                if 'afternoon' in routines:
                    p = doc.add_paragraph()
                    p.add_run('Afternoon:').bold = True
                    doc.add_paragraph(routines['afternoon'])
                
                if 'evening' in routines:
                    p = doc.add_paragraph()
                    p.add_run('Evening:').bold = True
                    doc.add_paragraph(routines['evening'])
        
        # Self-Care Strategies Section
        doc.add_heading('Self-Care Strategies', level=1)
        doc.add_paragraph('These strategies are tailored to support your grief healing journey.')
        
        self_care = guide.get('self_care', {})
        
        doc.add_heading('Physical Activity', level=2)
        doc.add_paragraph(self_care.get('physical_activity', ''))
        
        doc.add_heading('Nourishing Meal', level=2)
        doc.add_paragraph(self_care.get('nourishing_meal', ''))
        
        doc.add_heading('Evening Ritual', level=2)
        doc.add_paragraph(self_care.get('evening_ritual', ''))
        
        # Resources Section
        doc.add_heading('Resources', level=1)
        doc.add_paragraph('These resources have been curated to provide support for your specific situation.')
        
        resources = guide.get('resources', {})
        
        doc.add_heading('Books', level=2)
        doc.add_paragraph(resources.get('books', ''))
        
        doc.add_heading('Support Groups', level=2)
        doc.add_paragraph(resources.get('support_groups', ''))
        
        doc.add_heading('Online Resources', level=2)
        doc.add_paragraph(resources.get('online_resources', ''))
        
        # Generate DOCX as bytes
        docx_bytes = BytesIO()
        doc.save(docx_bytes)
        docx_bytes.seek(0)
        
        return docx_bytes
    
    except Exception as e:
        st.error(f"Error generating DOCX: {str(e)}")
        return None

def create_guide_txt(guide_data):
    """
    Create a TXT version of the guide
    
    Args:
        guide_data (dict): The guide data
        
    Returns:
        bytes: The TXT file as bytes
    """
    try:
        txt_output = BytesIO()
        
        # Title
        txt_output.write(f"{guide_data.get('title', 'Your Grief Guide')}\n".encode('utf-8'))
        txt_output.write(("=" * 50 + "\n\n").encode('utf-8'))
        
        # Introduction
        txt_output.write("INTRODUCTION\n".encode('utf-8'))
        txt_output.write(("-" * 20 + "\n").encode('utf-8'))
        txt_output.write(f"{guide_data.get('introduction', '')}\n\n".encode('utf-8'))
        
        # Weekly Routine
        txt_output.write("YOUR 7-DAY ROUTINE\n".encode('utf-8'))
        txt_output.write(("-" * 20 + "\n").encode('utf-8'))
        
        weekly_routine = guide_data.get("weekly_routine", {})
        for day in range(1, 8):
            day_str = str(day)
            if day_str in weekly_routine:
                txt_output.write(f"Day {day}:\n".encode('utf-8'))
                day_routine = weekly_routine[day_str]
                
                txt_output.write(f"  Morning: {day_routine.get('morning', '')}\n".encode('utf-8'))
                txt_output.write(f"  Afternoon: {day_routine.get('afternoon', '')}\n".encode('utf-8'))
                txt_output.write(f"  Evening: {day_routine.get('evening', '')}\n\n".encode('utf-8'))
        
        # Self-Care Strategies
        txt_output.write("SELF-CARE STRATEGIES\n".encode('utf-8'))
        txt_output.write(("-" * 20 + "\n").encode('utf-8'))
        
        self_care = guide_data.get("self_care", {})
        txt_output.write("Physical Activity:\n".encode('utf-8'))
        txt_output.write(f"  {self_care.get('physical_activity', '')}\n\n".encode('utf-8'))
        
        txt_output.write("Nourishing Meal:\n".encode('utf-8'))
        txt_output.write(f"  {self_care.get('nourishing_meal', '')}\n\n".encode('utf-8'))
        
        txt_output.write("Evening Ritual:\n".encode('utf-8'))
        txt_output.write(f"  {self_care.get('evening_ritual', '')}\n\n".encode('utf-8'))
        
        # Resources
        txt_output.write("HELPFUL RESOURCES\n".encode('utf-8'))
        txt_output.write(("-" * 20 + "\n").encode('utf-8'))
        
        resources = guide_data.get("resources", {})
        
        # Support Groups
        support_groups = resources.get("support_groups", [])
        if support_groups:
            txt_output.write("Support Groups:\n".encode('utf-8'))
            for group in support_groups:
                txt_output.write(f"  {group.get('name', '')}\n".encode('utf-8'))
                txt_output.write(f"  URL: {group.get('url', '')}\n".encode('utf-8'))
                txt_output.write(f"  {group.get('description', '')}\n\n".encode('utf-8'))
        
        # Books
        books = resources.get("books", [])
        if books:
            txt_output.write("Recommended Books:\n".encode('utf-8'))
            for book in books:
                txt_output.write(f"  {book.get('title', '')} by {book.get('author', '')}\n".encode('utf-8'))
                txt_output.write(f"  {book.get('description', '')}\n\n".encode('utf-8'))
        
        # Hotlines
        hotlines = resources.get("hotlines", [])
        if hotlines:
            txt_output.write("Hotlines:\n".encode('utf-8'))
            for hotline in hotlines:
                txt_output.write(f"  {hotline.get('name', '')}\n".encode('utf-8'))
                txt_output.write(f"  Number: {hotline.get('number', '')}\n".encode('utf-8'))
                txt_output.write(f"  {hotline.get('description', '')}\n\n".encode('utf-8'))
        
        # Professional Services
        prof_services = resources.get("professional_services", [])
        if prof_services:
            txt_output.write("Professional Services to Consider:\n".encode('utf-8'))
            for service in prof_services:
                txt_output.write(f"  {service.get('name', '')}\n".encode('utf-8'))
                txt_output.write(f"  {service.get('description', '')}\n\n".encode('utf-8'))
        
        txt_output.seek(0)
        return txt_output.getvalue()
        
    except Exception as e:
        st.error(f"Error creating TXT: {str(e)}")
        return None

def render_chat_widget(on_submit=None):
    """
    Render a chat widget for interacting with the guide
    
    Args:
        on_submit (callable): Function to call when the user submits a message
    """
    # Display chat history if available
    if 'guidance_conversation' in st.session_state and st.session_state.guidance_conversation:
        for entry in st.session_state.guidance_conversation:
            with st.chat_message("user"):
                st.write(entry.get('user_message', ''))
            with st.chat_message("assistant"):
                st.write(entry.get('ai_response', ''))
    
    # Chat input
    if user_input := st.chat_input("Ask a question about your grief guide..."):
        # Display user message
        with st.chat_message("user"):
            st.write(user_input)
        
        # Process and display response
        if on_submit:
            response = on_submit(user_input)
            
            with st.chat_message("assistant"):
                st.write(response)
            
            # Store conversation history
            if 'guidance_conversation' not in st.session_state:
                st.session_state.guidance_conversation = []
            
            st.session_state.guidance_conversation.append({
                'user_message': user_input,
                'ai_response': response
            })
            
            # Use rerun() to refresh UI with new chat history
            st.rerun() 